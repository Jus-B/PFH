import kivy
from docx import Document


# Imports
savedoc = input('Save the Document as: ')
input_doc = Document('Main_word.docx')
output_doc = Document()

docx1 = '.docx'
Savedfile = (savedoc + docx1)

def get_para_data(output_doc_name, paragraph):
    """
    Write the run to the new file and then set its font, bold, alignment, color etc. data.
    """

    output_para = output_doc_name.add_paragraph()
    for run in paragraph.runs:
        output_run = output_para.add_run(run.text)
        # Run's bold data
        output_run.bold = run.bold
        # Run's italic data
        output_run.italic = run.italic
        # Run's underline data
        output_run.underline = run.underline
        # Run's color data
        output_run.font.color.rgb = run.font.color.rgb
        # Run's font data
        output_run.style.name = run.style.name
    # Paragraph's alignment data
    output_para.paragraph_format.alignment = paragraph.paragraph_format.alignment

# Call the function
get_para_data(output_doc, input_doc.paragraphs[3])

# Save the new file
output_doc.save(Savedfile)
for para in input_doc.paragraphs:
    get_para_data(output_doc, para)

output_doc.save(Savedfile)


'''
from docx import Document

files = ['file1.docx', 'file2.docx']

def combine_word_documents(files):
    merged_document = Document()

    for index, file in enumerate(files):
        sub_doc = Document(file)

        # Don't add a page break if you've reached the last file.
        if index < len(files)-1:
           sub_doc.add_page_break()

        for element in sub_doc.element.body:
            merged_document.element.body.append(element)

    merged_document.save('merged.docx')

combine_word_documents(files)
    
    #pro Menü mit Checkbox if Checked add. else pass 
if true
    add_paragraph(text)
else: 
    false/pass
    
    '''
    